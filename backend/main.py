import io
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
import pdfplumber
from docx import Document

app = FastAPI(title="DerDieDas File Service")

SUPPORTED_FORMATS = [".pdf", ".docx", ".txt"]
MAX_FILE_SIZE = 10 * 1024 * 1024


def extract_text_from_pdf(file_content):
    all_text = []
    pdf_file = pdfplumber.open(io.BytesIO(file_content))
    
    for page in pdf_file.pages:
        page_text = page.extract_text()
        if page_text:
            all_text.append(page_text)
    
    pdf_file.close()
    return "\n".join(all_text)


def extract_text_from_docx(file_content):
    doc = Document(io.BytesIO(file_content))
    all_text = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            all_text.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            all_text.append(" | ".join(row_text))
    
    return "\n\n".join(all_text)


def extract_text_from_txt(file_content):
    encodings = ["utf-8", "latin-1", "cp1252"]
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            return text
        except UnicodeDecodeError:
            continue
    
    raise ValueError("Unsupported encoding")


def get_file_extension(filename):
    if not filename:
        return ""
    return Path(filename).suffix.lower()


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_extension = get_file_extension(file.filename)
    
    if file_extension not in SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format: {file_extension}, supported: {SUPPORTED_FORMATS}"
        )
    
    file_content = await file.read()
    file_size = len(file_content)
    
    if file_size == 0:
        raise HTTPException(status_code=400, detail="File is empty")
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large, max 10MB")
    
    try:
        if file_extension == ".pdf":
            text = extract_text_from_pdf(file_content)
        elif file_extension == ".docx":
            text = extract_text_from_docx(file_content)
        else:
            text = extract_text_from_txt(file_content)
        
        text = text.strip()
        
        if not text:
            raise HTTPException(status_code=422, detail="No text found in file")
        
        word_count = len(text.split())
        
        return {
            "success": True,
            "filename": file.filename,
            "text": text,
            "word_count": word_count
        }
        
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(status_code=422, detail=f"Error processing file: {str(error)}")
